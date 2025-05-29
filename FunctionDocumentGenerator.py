import os
import numpy as np
from qdrant_client import QdrantClient
from sentence_transformers import SentenceTransformer
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class FunctionDocumentGenerator:
    def __init__(self, qdrant_url, qdrant_api_key, collections=None):
        if collections is None:
            self.collection_names = ["Flexcube_user_guide_14.x", "Flexcube_Userguide_12.x"]
        else:
            self.collection_names = collections
        
        # Validate environment variables
        self.qdrant_url = qdrant_url or os.getenv("QDRANT_URL")
        self.qdrant_api_key = qdrant_api_key or os.getenv("QDRANT_API_KEY")
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        
        if not self.qdrant_url or not self.qdrant_api_key or not self.groq_api_key:
            raise ValueError("Missing required environment variables: QDRANT_URL, QDRANT_API_KEY, or GROQ_API_KEY")

        # Initialize Qdrant client
        self.qdrant_client = QdrantClient(
            url=self.qdrant_url,
            api_key=self.qdrant_api_key
        )
        
        # Initialize Groq client
        self.groq_client = Groq(api_key=self.groq_api_key)
        
        # Load a model that generates 896-dimensional vectors
        self.model = SentenceTransformer('all-MiniLM-L6-v2')

    def generate_embeddings(self, query):
        """
        Generate embeddings for the given query with padding to 896 dimensions
        """
        embeddings = self.model.encode(query)
        
        # Pad or truncate to 896 dimensions
        if embeddings.shape[0] < 896:
            padded_embedding = np.pad(embeddings, (0, 896 - embeddings.shape[0]), mode='constant')
        else:
            padded_embedding = embeddings[:896]
        
        return padded_embedding.tolist()

    def search_vector_db(self, query, top_k=5):
        """
        Search vector database for relevant information
        """
        query_embedding = self.generate_embeddings(query)
        results = []
        
        for collection in self.collection_names:
            try:
                search_result = self.qdrant_client.search(
                    collection_name=collection,
                    query_vector=query_embedding,
                    limit=top_k
                )
                results.extend(search_result)
            except Exception as e:
                print(f"Error searching collection {collection}: {e}")
        
        return results

    def generate_document_with_llama(self, function_requirement, context):
        """
        Generate document using LLaMA 3 from Groq API
        """
        prompt = f"""
        Generate a comprehensive function specification document based on the following requirements and context:

        Function Requirement: {function_requirement}
        Context from User Guide: {context}

        Create a detailed document with the following structure:
        1. *INTRODUCTION*  
        
           Brief overview of the document purpose and scope.
        
        2. *REQUIREMENT OVERVIEW*  
           Clear statement of the business requirements and objectives.
        
        3. *CURRENT FUNCTIONALITY*  
           Description of how the system currently handles these requirements.
        
        4. *PROPOSED FUNCTIONAL APPROACH*
           Detailed explanation of the proposed solution and implementation.
        
        The document should be technical, precise, and provide clear insights for implementation.
        Focus on these four main sections as they constitute the core content.
        
        Note: Additional sections like Validations, Interface Impact, Migration Impact, etc. will be included in the final document template but do not need to be generated.
        """

        response = self.groq_client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[
                {"role": "system", "content": "You are a technical documentation specialist."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )

        return response.choices[0].message.content

    def generate_function_document(self, function_requirement):
        """
        Main method to generate function-specific document
        """
        # Step 1: Search vector database for relevant context
        vector_search_results = self.search_vector_db(function_requirement)
        
        # Extract text from search results
        context = "\n".join([
            result.payload.get('text', '') 
            for result in vector_search_results 
            if result.payload and 'text' in result.payload
        ])
        
        # Step 2: Generate document using Llama 3
        generated_document = self.generate_document_with_llama(
            function_requirement, 
            context
        )
        
        return generated_document
    
    def add_bookmark(self, paragraph, bookmark_name):
        """
        Add a bookmark to a paragraph
        """
        run = paragraph.add_run()
        tag = run._r
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        tag.append(start)
        
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), bookmark_name)
        tag.append(end)

    def add_hyperlink(self, paragraph, text, bookmark_name):
        """
        Create a single-click hyperlink to a bookmark in the document that works reliably
        """
        from docx.oxml.shared import OxmlElement, qn
        from docx.opc.constants import RELATIONSHIP_TYPE
        
        # Create the w:hyperlink tag
        hyperlink = OxmlElement('w:hyperlink')
        
        # Set the w:anchor attribute to the bookmark name
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Set history to 1 (important for single-click)
        hyperlink.set(qn('w:history'), '1')
        
        # Create a new run element (w:r)
        new_run = OxmlElement('w:r')
        
        # Create run properties (rPr)
        rPr = OxmlElement('w:rPr')
        
        # Add color (blue)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Blue color
        rPr.append(color)
        
        # Add underline
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        
        # Add "hyperlink" style
        style = OxmlElement('w:rStyle')
        style.set(qn('w:val'), 'Hyperlink')
        rPr.append(style)
        
        # Add run properties to run
        new_run.append(rPr)
        
        # Add text
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        # Add run to hyperlink
        hyperlink.append(new_run)
        
        # Add the hyperlink to the paragraph
        paragraph._p.append(hyperlink)
        
        return hyperlink

    def add_dotted_toc_entry(self, doc, text, bookmark_name, page_number):
        """
        Add a Table of Contents entry with dotted line, single-click hyperlink, and page number
        """
        paragraph = doc.add_paragraph()
        
        # Create hyperlink to the bookmark
        self.add_hyperlink(paragraph, text, bookmark_name)
        
        # Calculate number of dots based on text length
        dots = "." * (60 - len(text))  # Approximate dots based on text length
        
        # Add dots
        paragraph.add_run(" " + dots + " ")
        
        # Add page number
        page_run = paragraph.add_run(str(page_number))
        page_run.bold = True  # Make page number bold
    
        return paragraph

    def create_table_manually(self, doc, rows, cols, headers=None):
        """
        Create a table manually to avoid index errors
        """
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Add headers if provided
        if headers:
            for i in range(min(len(headers), cols)):
                try:
                    cell = table.rows[0].cells[i]
                    cell.text = headers[i]
                    # Try to make headers bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                except Exception as e:
                    print(f"Warning: Could not set header {i}: {e}")
        
        return table

    def add_page_number(self, doc):
        """
        Add page numbers to the document footer
        """
        section = doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add page number field
        run = paragraph.add_run()
        field_code = OxmlElement('w:fldChar')
        field_code.set(qn('w:fldCharType'), 'begin')

        instruction_text = OxmlElement('w:instrText')
        instruction_text.set(qn('xml:space'), 'preserve')
        instruction_text.text = 'PAGE'

        field_char = OxmlElement('w:fldChar')
        field_char.set(qn('w:fldCharType'), 'end')

        run._r.append(field_code)
        run._r.append(instruction_text)
        run._r.append(field_char)

    def restart_page_numbering(self, doc):
        """
        Restart page numbering from 1 after the TOC
        """
        section_break = doc.add_section()
        section_break.start_type = 2  # New page
        section_break.different_first_page_header_footer = True
        
        # Set page number format to start at 1
        section_props = section_break._sectPr
        page_num_type = OxmlElement('w:pgNumType')
        page_num_type.set(qn('w:start'), '1')
        section_props.append(page_num_type)
    
    def save_as_word(self, text, logo_path=None, filename="function_specification.docx"):
        """
        Create a Word document from the generated text and return it as bytes
        """
        try:
            # Create a new Document
            doc = Document()
            
            # Add logo image if path is provided
            if logo_path:
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(1.77), height=Inches(1.02))
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                    # Add bold title text below the image
                    title_paragraph = doc.add_paragraph()
                    title_run = title_paragraph.add_run("Functional Specification Document\n [Bank Name]")
                    title_run.bold = True  # Make the text bold
                    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  
                except Exception as e:
                    print(f"Warning: Could not add logo: {e}")
            
            # Add Table of Contents heading
            toc_heading = doc.add_heading('Table of Contents', level=1)
            
            # Define TOC items and their corresponding bookmarks
            toc_items = [
                ("1. Introduction", "intro_bookmark"),
                ("2. Requirement Overview", "req_overview_bookmark"),
                ("3. Current Functionality", "current_func_bookmark"),
                ("4. Proposed Functionalty", "proposed_func_bookmark"),
                ("5. Validations and Error Messages", "validations_bookmark"),
                ("6. Interface Impact", "interface_bookmark"),
                ("7. Migration Impact", "migration_bookmark"),
                ("8. Assumptions", "assumptions_bookmark"),
                ("9. RS-FS Traceability", "traceability_bookmark"),
                ("10. Open and Closed Queries", "queries_bookmark"),
                ("11. Annexure", "annexure_bookmark")
            ]
            
            # Add TOC entries with hyperlinks
            for i, (item, bookmark) in enumerate(toc_items):
                self.add_dotted_toc_entry(doc, item, bookmark, i+1)
            
            # Insert a page break after the Table of Contents
            doc.add_page_break()
            
            # Add page numbers to the document
            self.add_page_number(doc)
            
            # Parse the generated text and add core sections
            sections = {
                "1. INTRODUCTION": "",
                "2. REQUIREMENT OVERVIEW": "",
                "3. CURRENT FUNCTIONALITY": "",
                "4. PROPOSED FUNCTIONAL APPROACH": ""
            }
            
            # Process the generated text to extract sections
            current_section = None
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                # Check if line is a section header
                for section_title in sections.keys():
                    # Flexible matching for section titles
                    if any(title.lower() in line.lower() for title in [
                        section_title.lower(), 
                        section_title.split(". ")[1].lower(),
                        section_title.split(". ")[0] + "." + section_title.split(". ")[1].lower()
                    ]):
                        current_section = section_title
                        break
            
                if current_section and line.lower() not in [s.lower() for s in sections.keys()]:
                    if not any(title.lower() in line.lower() for title in [t.split(". ")[1].lower() for t in sections.keys()]):
                        sections[current_section] += line + "\n"
            
            # Map section titles to bookmark names
            section_bookmarks = {
                "1. INTRODUCTION": "intro_bookmark",
                "2. REQUIREMENT OVERVIEW": "req_overview_bookmark",
                "3. CURRENT FUNCTIONALITY": "current_func_bookmark",
                "4. PROPOSED FUNCTIONAL APPROACH": "proposed_func_bookmark"
            }
            
            # Add each section to the document with bookmarks
            for i, (section_title, content) in enumerate(sections.items(), 1):
                # Create a heading
                heading = doc.add_heading(f"{i}. {section_title.split('. ')[1].title()}", level=1)
                
                # Add bookmark to this heading
                self.add_bookmark(heading, section_bookmarks[section_title])
                
                # Add content
                if content.strip():
                    paragraph = doc.add_paragraph(content)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Justify the text
                    # Set paragraph spacing
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(6)
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15  # Adjust line spacing
                else:
                    paragraph = doc.add_paragraph("Content to be added.")
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Justify the text
                    # Set paragraph spacing
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(6)
                    paragraph_format.space_after = Pt(6)
                    paragraph_format.line_spacing = 1.15  # Adjust line spacing
            
            # Add additional template sections with bookmarks
            additional_sections = [
                ("5. Validations and Error Messages", "NA.", "validations_bookmark"),
                ("6. Interface Impact", "NA.", "interface_bookmark"),
                ("7. Migration Impact", "NA", "migration_bookmark"),
                ("8. Assumptions", "To be determined.", "assumptions_bookmark"),
                ("11. Annexure", "To be added as required.", "annexure_bookmark")
            ]
            
            for title, content, bookmark in additional_sections:
                # Create a heading
                heading = doc.add_heading(title, level=1)
                
                # Add bookmark to this heading
                self.add_bookmark(heading, bookmark)
                
                # Add content
                paragraph = doc.add_paragraph(content)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Justify the text
                # Set paragraph spacing
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(6)
                paragraph_format.space_after = Pt(6)
                paragraph_format.line_spacing = 1.15  # Adjust line spacing
            
            # Add RS-FS Traceability section with table and bookmark
            heading = doc.add_heading("9. RS-FS Traceability", level=1)
            self.add_bookmark(heading, "traceability_bookmark")
            
            # Create simple table for RS-FS Traceability
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            
            # Add headers manually
            try:
                headers = ["S. No.", "RS Section", "RS Section Description", "FS Section / Description"]
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
            except Exception as e:
                print(f"Warning: Issue with traceability table: {e}")
            
            # Add Open and Closed Queries section with table and bookmark
            heading = doc.add_heading("10. Open and Closed Queries", level=1)
            self.add_bookmark(heading, "queries_bookmark")
            
            # Create queries table with simple structure
            query_table = doc.add_table(rows=2, cols=6)
            query_table.style = 'Table Grid'
            
            # Add headers manually
            try:
                query_headers = ["Sr. No", "Issue Details", "Date Raised", "Clarification", "Raised By", "Current Status"]
                for i, header in enumerate(query_headers):
                    query_table.rows[0].cells[i].text = header
            except Exception as e:
                print(f"Warning: Issue with queries table: {e}")
            
            # Return the document object
            return doc
        except Exception as e:
            print(f"Error generating Word document: {e}")
            return None